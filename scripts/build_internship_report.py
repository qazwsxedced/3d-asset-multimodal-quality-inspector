from __future__ import annotations

import math
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "report_assets"
DOCX_PATH = OUT_DIR / "3d_asset_multimodal_quality_inspection_internship_report_zh.docx"

ARCHITECTURE = OUT_DIR / "3d_asset_multimodal_quality_inspector" / "slide-3.png"
BLENDER_DIR = ROOT / "data" / "blender_research_v3" / "images" / "asset_00000"
SCREEN_RULE = Path(r"C:\Users\Cr7\xwechat_files\wxid_hux8d2iznbwv22_0e31\temp\RWTemp\2026-08\9e20f478899dc29eb19741386f9343c8\03dc0c31ce2215793fdfd09d9161b4d2.png")
SCREEN_VLM = Path(r"C:\Users\Cr7\xwechat_files\wxid_hux8d2iznbwv22_0e31\temp\RWTemp\2026-08\9e20f478899dc29eb19741386f9343c8\2cfc683fc3e4aebd3f9f29ce0d1accd6.png")
SCREEN_HYBRID = Path(r"C:\Users\Cr7\xwechat_files\wxid_hux8d2iznbwv22_0e31\temp\RWTemp\2026-08\9e20f478899dc29eb19741386f9343c8\8539892ccb78028b0a483791abe150fd.png")
SCREEN_PASS = Path(r"C:\Users\Cr7\xwechat_files\wxid_hux8d2iznbwv22_0e31\temp\RWTemp\2026-08\9e20f478899dc29eb19741386f9343c8\22922a8fbdf19424a066da20c079ee5c.png")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1F6B4F"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def font_path(name: str) -> str:
    candidates = [Path(r"C:\Windows\Fonts") / name, Path(r"C:\Windows\Fonts\arial.ttf")]
    for path in candidates:
        if path.exists():
            return str(path)
    return str(candidates[-1])


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path("arialbd.ttf" if bold else "arial.ttf"), size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    widths_dxa = [int(round(w * 1440)) for w in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Microsoft YaHei", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True


def add_text(doc, text, style=None, bold=False, color=INK, size=11, align=None, before=0, after=6, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=1.10)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_paragraph(doc, parts, style=None, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=1.10)
    for text, kwargs in parts:
        run = p.add_run(text)
        set_run_font(run, **kwargs)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.50 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, after=5, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.50)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, after=5, line=1.167)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    color = BLUE if level < 3 else DARK_BLUE
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_run_font(run, size=size, color=color, bold=True)
    set_paragraph_spacing(p, before=16 if level == 1 else 12 if level == 2 else 8, after=8 if level == 1 else 6 if level == 2 else 4, line=1.0, keep=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=8, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def add_picture(doc, path: Path, width=6.20, caption=None):
    if not path.exists():
        add_text(doc, f"[图片缺失：{path.name}]", color=RED, size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=2, line=1.0)
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    alt_text = caption or path.stem
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", alt_text)
    doc_pr.set("descr", alt_text)
    if caption:
        add_caption(doc, caption)


def add_callout(doc, label, text, fill=CALLOUT, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5], indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=3, line=1.10)
    label_run = p.add_run(label + "  ")
    set_run_font(label_run, size=10.5, color=label_color, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, font_size=9.4, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0, line=1.0)
        run = p.add_run(str(value))
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=0, line=1.05)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, code, caption=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5], indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "17202A")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    run = p.add_run(code.strip())
    set_run_font(run, name="Consolas", size=8.3, color="E5E7EB")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if caption:
        add_caption(doc, caption)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Bullet 2", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167


def make_evidence_collage(path: Path) -> None:
    images = [
        (BLENDER_DIR / "view_0.png", "View 0"),
        (BLENDER_DIR / "view_1.png", "View 1"),
        (BLENDER_DIR / "view_2.png", "View 2"),
        (BLENDER_DIR / "view_3.png", "View 3"),
        (BLENDER_DIR / "uv.png", "UV diagnostic"),
        (BLENDER_DIR / "normal.png", "Normal diagnostic"),
    ]
    canvas = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(34, True)
    label_font = pil_font(22, True)
    draw.text((70, 34), "Blender multimodal evidence", fill="#0B2545", font=title_font)
    tile_w, tile_h = 520, 440
    x_positions = [70, 640, 1210]
    y_positions = [110, 650]
    for idx, (image_path, label) in enumerate(images):
        if not image_path.exists():
            continue
        img = Image.open(image_path).convert("RGB")
        img.thumbnail((tile_w, tile_h - 42))
        x = x_positions[idx % 3]
        y = y_positions[idx // 3]
        box = Image.new("RGB", (tile_w, tile_h), "#F2F4F7")
        bx = (tile_w - img.width) // 2
        by = 38 + (tile_h - 42 - img.height) // 2
        box.paste(img, (bx, by))
        box_draw = ImageDraw.Draw(box)
        box_draw.text((18, 10), label, fill="#1F4D78", font=label_font)
        canvas.paste(box, (x, y))
    canvas.save(path)


def draw_bar(draw, x, y, width, height, value, color, label, max_value=100):
    draw.rounded_rectangle((x, y, x + width, y + height), radius=8, fill="#E9EEF4")
    filled = int(width * value / max_value)
    draw.rounded_rectangle((x, y, x + filled, y + height), radius=8, fill=color)
    draw.text((x + width + 14, y - 2), f"{value:.2f}%", fill="#1F2937", font=pil_font(21, True))
    draw.text((x - 150, y - 2), label, fill="#344054", font=pil_font(21))


def make_results_chart(path: Path) -> None:
    canvas = Image.new("RGB", (1900, 1120), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 38), "Experimental results", fill="#0B2545", font=pil_font(40, True))
    draw.text((70, 88), "Percentages; left: 120-sample Blender test, right: 28 controlled .blend fixtures", fill="#667085", font=pil_font(22))
    sections = [
        ("Phase 2 B0-B4", [
            ("Rule", 95.00, 95.71),
            ("B0", 69.17, 0.00),
            ("B1", 44.17, 0.00),
            ("B2", 20.00, 0.00),
            ("B3", 32.50, 0.00),
            ("B4", 88.33, 82.98),
        ]),
        ("External mode comparison", [
            ("Rule-only", 92.86, 94.44),
            ("VLM-only", 89.29, 90.00),
            ("Hybrid", 92.86, 94.44),
        ]),
    ]
    panel_x = [70, 1010]
    panel_w = 800
    for panel_idx, (title, rows) in enumerate(sections):
        x0 = panel_x[panel_idx]
        draw.rounded_rectangle((x0, 160, x0 + panel_w, 1020), radius=18, outline="#D0D5DD", width=3, fill="#FBFCFD")
        draw.text((x0 + 32, 188), title, fill="#1F4D78", font=pil_font(28, True))
        draw.text((x0 + 32, 235), "Blue: quality accuracy   Orange: defect Macro-F1", fill="#667085", font=pil_font(18))
        y = 300
        for label, quality, f1 in rows:
            draw.text((x0 + 32, y - 3), label, fill="#344054", font=pil_font(21, True))
            bar_x = x0 + 170
            draw_bar(draw, bar_x, y, 250, 27, quality, "#2E74B5", "", 100)
            draw_bar(draw, bar_x + 325, y, 250, 27, f1, "#D97706", "", 100)
            y += 105 if len(rows) <= 3 else 110
    draw.text((70, 1060), "Interpretation: B4 substantially improves task-specific VLM output; the deterministic rule remains the hard safety gate in the prototype.", fill="#667085", font=pil_font(20))
    canvas.save(path)


def add_cover(doc):
    add_text(doc, "项目研究与实习总结报告", color=BLUE, size=11, bold=True, after=8)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=8, line=1.0)
    run = p.add_run("面向多图与结构化元数据融合的\n视觉语言模型推理研究")
    set_run_font(run, size=25, color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=18, line=1.15)
    r = p2.add_run("3D 资产多模态质量检验系统的研究、实现与工业化思考")
    set_run_font(r, size=14, color=MUTED)
    meta = [
        ("研究对象", "多视角图像、UV/法线诊断图、低层结构化几何元数据与 VLM 推理"),
        ("核心模型", "Qwen2.5-VL-3B-Instruct + QLoRA/PEFT"),
        ("工程栈", "Blender 5.2、Python、PyTorch、Transformers、Gradio、HTML/JSON 审计"),
        ("硬件环境", "NVIDIA GeForce RTX 4060 Laptop GPU，8 GB VRAM"),
        ("报告定位", "实习项目总结 / 研究过程记录 / 工业化原型说明"),
    ]
    table = doc.add_table(rows=0, cols=2)
    for label, value in meta:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for idx, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.05)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, color=NAVY if idx == 0 else INK, bold=idx == 0)
    set_table_geometry(table, [1.35, 5.15])
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(doc, "一句话概括", "我没有把项目包装成通用 VLM 推理，而是围绕 3D 资产质检这一具体工业任务，研究多图与结构化元数据如何共同影响结构化诊断，并把研究结果落成可上传 .blend、可复核、可导出审计报告的原型系统。", fill=CALLOUT)
    add_picture(doc, ASSET_DIR / "evidence_collage.png", width=6.20, caption="图 1  项目输入从普通渲染扩展到多视角、UV 与法线诊断证据")
    add_text(doc, "本报告根据项目代码、实验报告、GPU 实际运行记录和 Demo 验证截图整理。数字均来自当前已完成的受控实验；真实客户资产与生产 SLA 仍属于后续工作。", color=MUTED, size=9.5, italic=True, after=0)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_evidence_collage(ASSET_DIR / "evidence_collage.png")
    make_results_chart(ASSET_DIR / "results_chart.png")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    set_paragraph_spacing(header, before=0, after=0, line=1.0)
    left = header.add_run("3D Asset Multimodal Quality Inspector")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    tab = header.add_run("\t研究与实习总结报告")
    set_run_font(tab, size=8.5, color=MUTED)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    footer = section.footer.paragraphs[0]
    set_page_number(footer)

    add_cover(doc)
    add_page_break(doc)

    add_heading(doc, "摘要与核心结论", 1)
    add_text(doc, "本项目面向游戏、影视渲染、数字孪生和内容生产中的 3D 资产质量检查问题，研究多视角图像、UV/法线诊断图与结构化几何元数据对视觉语言模型（VLM）结构化推理能力的影响。研究没有把目标表述为“通用 VLM 推理”，而是限定为“面向多图与结构化元数据融合的视觉语言模型推理研究”。", after=8)
    add_text(doc, "项目完成了从 Blender 数据生成、B0-B4 受控实验、Qwen2.5-VL QLoRA 微调、结构化 JSON 评测，到 .blend 上传、规则/VLM/Hybrid 审核和 HTML 审计报告的完整链路。600 样本多资产测试中，B4 单次实验达到 88.33% 质量准确率、82.98% 缺陷 Macro-F1 和 99.17% Schema 有效率；三随机种子下缺陷 Macro-F1 为 82.55% ± 0.72%。", after=8)
    add_callout(doc, "核心判断", "VLM 适合承担多模态解释、结构化输出和修复建议；确定性的几何规则适合承担硬质量门禁。两者冲突时进入 REVIEW REQUIRED，是当前原型最重要的工业安全边界。", fill="EEF6FF", label_color=BLUE)
    add_heading(doc, "报告结构", 2)
    for item in [
        "行业不足与研究问题：为什么 3D 质检值得做，以及本项目针对的具体缺口。",
        "研究思路与具体做法：数据、实验、模型、评测和工程系统如何连接。",
        "过程与证据：从环境搭建、Blender 生成、GPU 训练到 Demo 验证的完成过程。",
        "实验结果与边界：主实验、三种策略对比、外部输入验证和错误分析。",
        "行业看法、未来规划与实习总结：项目能落到哪里，还需要补齐什么。",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "一、行业背景：3D 质量检查仍存在什么不足？", 1)
    add_text(doc, "3D 资产生产链路已经高度数字化，但质量检查仍常常依赖技术美术、模型工程师或内容审核人员逐项打开 DCC 软件检查。问题不在于“有没有检查工具”，而在于检查结果分散、标准不统一、复杂证据难以复核，以及生成式模型缺少安全边界。", after=8)
    rows = [
        ("人工检查碎片化", "拓扑、UV、法线、破洞、退化面等问题分布在不同视图和工具中", "检查耗时，经验依赖强，难以稳定复现"),
        ("视觉与结构信息割裂", "普通渲染能看外观，但无法完整表达 UV 重叠、非流形和退化面", "多模态证据没有形成统一诊断输入"),
        ("VLM 输出不稳定", "模型可能能生成 JSON，却把缺陷名称、严重程度或字段含义混淆", "不能直接作为自动放行的硬门禁"),
        ("缺少可审计流程", "很多 Demo 只展示一次推理结果，没有保存输入、版本、规则和冲突原因", "出现误判后难以追溯与复核"),
        ("实验与工业脱节", "合成 benchmark 容易只测模型分数，工程系统又常缺乏严谨对照", "研究结果难转化为可运行的生产流程"),
    ]
    add_table(doc, ["行业不足", "具体表现", "直接后果"], rows, [1.35, 2.55, 2.60], font_size=9.2)
    add_heading(doc, "我针对的是哪个问题？", 2)
    add_text(doc, "我的研究针对的是一个更具体、可验证的问题：在 3D 资产质量诊断中，多视角图像、UV/法线诊断图和低层结构化几何统计分别以及联合使用时，会对 VLM 的缺陷识别、质量判断、严重程度预测和 JSON 结构化输出产生多大影响？", after=6)
    add_text(doc, "因此，项目不以“让模型看起来很聪明”为目标，而以以下四个可验证目标为中心：", after=4)
    for item in [
        "建立严格的 B0-B4 输入协议，隔离每一种模态的贡献。",
        "建立字段级评测，区分 JSON 格式正确、质量判断正确和缺陷集合正确。",
        "建立 Rule-only、VLM-only、Hybrid 三种策略的同输入对照。",
        "把研究结果封装为可上传 .blend、可生成证据、可触发人工复核的本地原型。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、研究思路：从“单点模型”转为“证据链系统”", 1)
    add_text(doc, "项目的总体思路是把 3D 质检拆成四层：输入层生成一致的多模态证据，规则层负责确定性安全边界，VLM 层负责复杂解释和结构化建议，审核层负责把冲突变成可追踪的人工复核任务。", after=8)
    add_picture(doc, ARCHITECTURE, width=6.20, caption="图 2  系统整体链路：.blend 上传 → Blender 预处理 → 多模态证据 → Rule + VLM → JSON/HTML 审计")
    add_heading(doc, "研究假设", 2)
    for item in [
        "多视角图像能够减少单视角遮挡，提高几何外观判断的完整性。",
        "UV/法线图能为普通渲染难以表达的问题提供诊断证据，但需要任务相关训练才能被 VLM 稳定利用。",
        "结构化几何统计对规则系统非常有效，但直接拼入 prompt 不等于模型能够可靠理解其诊断含义。",
        "规则与 VLM 的组合不应简单投票，而应以规则作为质量门禁，以不一致检测触发人工复核。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "系统安全策略", 2)
    add_callout(doc, "安全边界", "当规则与 VLM 在质量、缺陷集合、严重程度或 Schema 有效率上不一致时，系统不自动放行，而是输出 REVIEW REQUIRED，同时保存双方结果与冲突原因。", fill="FFF8E8", label_color=GOLD)

    add_page_break(doc)
    add_heading(doc, "三、数据与实验设计：如何把问题变成可研究的实验？", 1)
    add_heading(doc, "1. Blender 多资产数据生成", 2)
    add_text(doc, "我使用 Blender 自动生成可控的 3D 资产与质量问题，覆盖 ico_sphere、cube、cylinder、cone、torus 五类资产，并注入六类缺陷：non-manifold、UV overlap、flipped normals、hole、stretched triangles、degenerate faces。每个样本保存四个相机视角、UV 诊断图、法线诊断图、低层几何统计和结构化问答。", after=8)
    add_picture(doc, ASSET_DIR / "evidence_collage.png", width=6.20, caption="图 3  同一资产的四视角、UV 诊断图与法线诊断图；所有输入均由 Blender 管线生成")
    add_table(doc, ["数据项", "当前实现"], [
        ("主 benchmark", "600 samples；5 个资产族；420 / 60 / 120 train / val / test"),
        ("缺陷类型", "6 类受控缺陷；标签来自 Blender 注入逻辑与几何统计"),
        ("切分方式", "按 scene_id 分组切分，避免同场景泄漏到不同 split"),
        ("外部输入验证", "28 个 controlled .blend fixtures；28/28 预处理成功"),
        ("数据边界", "属于受控合成/外部输入验证，不等同于真实客户人工标注数据"),
    ], [1.55, 4.95], font_size=9.3)
    add_heading(doc, "2. B0-B4 受控实验协议", 2)
    add_text(doc, "核心原则是固定 sample_id、固定测试集和固定评测器，只逐步改变输入条件。这样可以把“多图有没有用”“诊断图有没有用”“元数据有没有用”“LoRA 是否提升”分开回答。", after=6)
    add_table(doc, ["条件", "输入", "研究目的"], [
        ("B0", "单张普通渲染图", "单图零样本基线"),
        ("B1", "四张不同视角渲染图", "隔离多视角信息贡献"),
        ("B2", "多视角 + UV/法线图", "隔离诊断图贡献"),
        ("B3", "多视角 + 结构化几何元数据", "隔离低层统计信号贡献"),
        ("B4", "B3 + QLoRA/SFT", "验证任务适配后的结构化推理能力"),
    ], [0.70, 2.45, 3.35], font_size=9.2)
    add_heading(doc, "3. 评测指标", 2)
    for item in [
        "JSON valid rate：输出能否被解析为 JSON。",
        "Schema valid rate：字段、枚举值和列表类型是否满足协议。",
        "Quality / severity accuracy：质量结论和严重程度的字段级准确率。",
        "Defect Macro-F1：多标签缺陷集合的均衡指标，避免只看 clean/pass。",
        "Field exact accuracy：缺陷集合、修复计划等字段是否完全一致。",
        "Latency P50/P95：面向实际批处理与审核链路的耗时指标。",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "四、具体实现：我做了哪些代码和工程工作？", 1)
    add_heading(doc, "1. Blender 预处理与几何统计", 2)
    add_text(doc, "Blender 脚本负责打开 .blend、隔离任务目录、生成四个视角、UV/法线诊断图，并从 mesh 中提取 vertex_count、face_count、boundary_edge_count、non_manifold_edge_count、flipped_normal_count、degenerate_face_count、uv_overlap_ratio 和 triangle_area_stats 等字段。低层统计只作为证据输入，不直接把 is_defect 这类答案字段塞给模型。", after=8)
    add_code_block(doc, '''# scripts/run_external_batch.py：每个 .blend 独立目录处理并记录重试\nwhile attempts <= args.retries:\n    attempts += 1\n    completed = subprocess.run(command, capture_output=True,\n                               text=True, timeout=args.timeout)\n    if completed.returncode == 0:\n        error = None\n        break\n    error = (completed.stderr or completed.stdout)[-2000:]''', "代码节选 1  外部 .blend 批处理：隔离目录、失败重试、错误尾部记录")
    add_heading(doc, "2. Qwen2.5-VL 与 QLoRA 适配", 2)
    add_text(doc, "模型使用 Qwen2.5-VL-3B-Instruct，训练采用 4-bit QLoRA 以适配 8 GB 显存。训练入口支持多图输入、结构化 prompt、LoRA rank 16、alpha 32、dropout 0.05、学习率 2e-4、梯度累积 8 和 1 epoch。推理端支持 offload_dir，避免 8 GB 显存下加载适配器时出现缓冲区不足。", after=8)
    add_table(doc, ["配置项", "取值 / 说明"], [
        ("Base model", "Qwen/Qwen2.5-VL-3B-Instruct"),
        ("Quantization", "4-bit loading；NF4 compute path"),
        ("LoRA", "rank=16，alpha=32，dropout=0.05"),
        ("Training", "1 epoch，gradient accumulation=8，learning rate=2e-4"),
        ("Image budget", "约 50,176—100,352 pixels，面向 8 GB GPU 做约束"),
        ("Output", "单一 JSON 对象；quality / defect_types / severity / repair_plan"),
    ], [1.55, 4.95], font_size=9.3)
    add_heading(doc, "3. Rule / VLM / Hybrid 对比", 2)
    add_text(doc, "Rule-only 只使用结构化统计和训练集拟合出的稳健阈值；VLM-only 只使用 VLM 诊断结果；Hybrid 只有在 VLM Schema 有效且与规则在质量、严重程度和缺陷集合上达成一致时才采纳 VLM，否则选择规则结果并标记人工复核。", after=8)
    add_code_block(doc, '''# scripts/compare_modes.py：保守 Hybrid 的核心决策\nif schema_valid and quality_agrees and severity_agrees and defects_agree:\n    selected_source = "vlm"\n    review_required = False\nelse:\n    selected_source = "rule_gate_plus_vlm_explanation"\n    review_required = True''', "代码节选 2  Hybrid 审核：不一致不自动放行")

    add_heading(doc, "五、完成过程：项目是怎样一步步落地的？", 1)
    add_table(doc, ["阶段", "完成内容", "形成的证据"], [
        ("阶段 1：协议与 mock", "完成 B0-B4 schema、数据生成器、校验器、JSON 评测器", "120 条 mock 样本；84/12/24 split；数据 errors=0"),
        ("阶段 2：GPU 环境", "解决 Python、PyTorch CUDA、Transformers、PEFT、bitsandbytes 与 Qwen 权重加载", "RTX 4060 8 GB；4-bit load OK；GPU smoke test OK"),
        ("阶段 3：真实 VLM", "完成 B0-B3 zero-shot 与 B4 QLoRA 训练、推理和指标汇总", "B0-B4 结果表；三 seed 稳定性；错误分析"),
        ("阶段 4：Blender 真实管线", "修复 Blender 5.2 引擎兼容、UV/法线诊断渲染和 .blend 资产预处理", "多视角、UV、法线图；30 样本校准；600 样本 benchmark"),
        ("阶段 5：工程原型", "完成上传、重试、批处理、审计 JSON/HTML 和 Rule/VLM/Hybrid 三模式", "Gradio Demo；28 个外部 fixture 28/28 成功"),
        ("阶段 6：公开交付", "整理 README、技术报告、论文中英文版、PPT、简历表述并完成 GitHub 清理", "release/public-research 分支与脱敏 HTML 证据"),
    ], [1.25, 3.25, 2.00], font_size=8.9)
    add_heading(doc, "过程中的关键问题与解决", 2)
    for item in [
        "Blender 5.2 不再识别旧的 BLENDER_EEVEE_NEXT 枚举：改为兼容当前版本的渲染引擎选择。",
        "UV/法线诊断材质节点接口在 Blender 5.2 发生差异：调整 Combine/Shader 节点连接，并补充无相机保护。",
        "pip 下载 PyTorch wheel 出现 hash mismatch：改用 curl 下载 wheel 后本地安装，完成 CUDA 验证。",
        "Transformers 版本对 TrainingArguments 参数存在差异：移除不兼容参数后完成 dry-run 和全量训练。",
        "8 GB 显存加载 LoRA 适配器时需要 offload_dir：增加推理 offload 参数，完成完整测试集推理。",
        "初始任务语义存在空间问答与质量诊断漂移：重构为质量/缺陷/严重程度/修复计划 schema，避免 metadata 直接泄漏答案。",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "六、实验结果：这种做法带来了什么提升？", 1)
    add_text(doc, "主实验使用 600 样本多资产 Blender benchmark，测试集为 120 条。以下结果来自最终多资产报告；Rule 是确定性元数据基线，B0-B3 是 zero-shot 输入条件，B4 是加入 QLoRA 后的任务适配结果。", after=8)
    add_picture(doc, ASSET_DIR / "results_chart.png", width=6.20, caption="图 4  主实验与外部 .blend 验证的指标对比；图中数值为当前受控实验结果")
    add_table(doc, ["条件", "Quality", "Severity", "Defect Macro-F1", "Schema valid", "Mean latency"], [
        ("Rule baseline", "95.00%", "93.33%", "95.71%", "100.00%", "规则推理"),
        ("B0", "69.17%", "48.33%", "0.00%", "0.00%", "4.62 s"),
        ("B1", "44.17%", "51.67%", "0.00%", "0.00%", "4.89 s"),
        ("B2", "20.00%", "53.33%", "0.00%", "0.00%", "4.93 s"),
        ("B3", "32.50%", "43.33%", "0.00%", "0.00%", "4.16 s"),
        ("B4 + LoRA", "88.33%", "81.67%", "82.98%", "99.17%", "4.71 s"),
    ], [1.25, 0.90, 0.90, 1.25, 1.10, 1.10], font_size=8.8)
    add_callout(doc, "结果解释", "B4 相比 B0 的质量准确率提升 19.16 个百分点，缺陷 Macro-F1 从当前协议下的 0 提升到 82.98%。这说明任务适配对结构化输出和缺陷集合学习有效；但在外部受控 .blend 上，Rule-only 仍略高于 VLM-only，因此不能把 VLM 宣称为规则系统的替代品。", fill="EEF6FF", label_color=BLUE)
    add_heading(doc, "三随机种子稳定性", 2)
    add_table(doc, ["指标", "均值 ± 样本标准差"], [
        ("Quality accuracy", "85.83% ± 4.33%"),
        ("Severity accuracy", "75.56% ± 6.25%"),
        ("Defect Macro-F1", "82.55% ± 0.72%"),
        ("JSON / Schema valid rate", "94.72% ± 7.70%"),
        ("Unseen scene defect Macro-F1", "84.19% ± 1.48%"),
        ("Unseen question type defect Macro-F1", "52.19% ± 12.55%"),
    ], [3.25, 3.25], font_size=9.5)
    add_heading(doc, "28 个外部 .blend 受控输入验证", 2)
    add_table(doc, ["策略", "Quality", "Severity", "Defect Macro-F1", "Review rate"], [
        ("Rule-only", "92.86%", "92.86%", "94.44%", "0.00%"),
        ("VLM-only", "89.29%", "89.29%", "90.00%", "0.00%"),
        ("Hybrid", "92.86%", "92.86%", "94.44%", "10.71%"),
    ], [1.70, 1.15, 1.15, 1.45, 1.05], font_size=9.2)
    add_text(doc, "预处理方面，28 个资产全部成功，平均耗时 3.50 s，P50 为 3.41 s，P95 为 3.83 s；25/28 个样本三种策略达成一致，3/28 个样本被路由为人工复核。该结果证明工程链路可运行，但数据仍是 Blender 受控生成的 fixture，不是客户人工标注 benchmark。", after=6)

    add_page_break(doc)
    add_heading(doc, "七、网页 Demo 与过程截图", 1)
    add_text(doc, "为了证明项目不仅停留在离线评测，我实现了一个本地 Gradio 原型：用户上传 .blend，系统调用 Blender 后台预处理，展示四视角、UV/法线证据和结构化 metadata，并允许切换 Rule baseline、VLM diagnosis、Hybrid review，最后导出 JSON/HTML 审计报告。", after=8)
    add_picture(doc, SCREEN_RULE, width=6.20, caption="图 5  上传 UV-overlap 资产后，Rule baseline 输出 FAIL，并识别 stretched_triangles 与 uv_overlap")
    add_picture(doc, SCREEN_VLM, width=6.20, caption="图 6  VLM diagnosis 输出 uv_overlap / medium；与规则结果不一致，因此进入复核逻辑")
    add_picture(doc, SCREEN_HYBRID, width=6.20, caption="图 7  Hybrid review 保存规则结果、VLM 结果和 disagreement reasons，并输出 REVIEW REQUIRED")
    add_picture(doc, SCREEN_PASS, width=6.20, caption="图 8  clean .blend 资产的通过路径：PASS、空缺陷集合、severity=none")
    add_callout(doc, "工业化意义", "Demo 的价值不在于把一个模型结果做成网页，而在于把“证据生成—规则门禁—模型解释—冲突复核—审计导出”连接成一个可以被技术美术或审核人员理解的操作闭环。", fill="F4F6F9", label_color=DARK_BLUE)

    add_page_break(doc)
    add_heading(doc, "八、行业看法：我认为这个方向真正有价值的地方", 1)
    add_text(doc, "我对行业的判断是：3D 质检不会简单地被一个 VLM 取代，也不应该把“生成式模型能看图”直接等同于“模型可以自动放行资产”。真正有价值的方向，是把 AI 放到人工审核最耗时、最需要跨证据解释的环节，并保留可验证的安全边界。", after=8)
    add_heading(doc, "我认为可以产生价值的三类场景", 2)
    for item in [
        "游戏与影视资产生产：批量检查拓扑、UV、法线和三角形质量，减少技术美术逐个打开资产的时间。",
        "数字孪生与仿真资产：在进入仿真或渲染前，自动检查资产是否满足几何和可渲染性要求，并保存版本化审计记录。",
        "内容平台与资产服务器：将 .blend、FBX、GLB 等格式接入批处理队列，对失败资产给出缺陷类型、严重程度和修复建议。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "对未来产品化的期待", 2)
    add_table(doc, ["方向", "下一步补充", "工业价值"], [
        ("数据", "客户资产、人工复核标签、跨团队缺陷本体", "提高分布覆盖和可信度"),
        ("模型", "置信度校准、缺陷定位、视觉证据引用、轻量模型蒸馏", "降低误报漏报并便于部署"),
        ("系统", "任务队列、并发控制、权限、持久化、监控和告警", "从本地 Demo 走向服务化"),
        ("集成", "Blender/Maya 插件、资产服务器 API、CI/CD 质量门禁", "嵌入现有生产流程"),
        ("评测", "真实人工标注集、跨团队复核一致性、吞吐和 SLA", "形成可被业务采用的指标体系"),
    ], [1.15, 3.15, 2.20], font_size=9.1)
    add_callout(doc, "边界声明", "当前项目证明了研究管线、工程原型和审核边界的可行性，但没有证明跨行业泛化、真实客户分布上的可靠性，也没有达到生产部署标准。", fill="FFF1F2", label_color=RED)

    add_heading(doc, "九、未来扩展计划", 1)
    for item in [
        "扩展真实或人工标注的 20—50 个外部资产集，并记录标签来源、复核者和一致性。",
        "增加针对 stretched triangles、破洞和复杂 UV 岛的困难样本与缺陷定位标注。",
        "增加多次随机种子、置信度校准、阈值选择和成本敏感评测。",
        "完成批量队列、失败重试、并发限流、日志检索和持久化审计数据库。",
        "将输入扩展到 FBX/GLB，并通过 Blender/Maya 插件或资产服务器 API 接入生产链路。",
        "评测本地推理与小模型蒸馏后的吞吐、显存占用和 P95 端到端延迟。",
    ]:
        add_number(doc, item)

    add_page_break(doc)
    add_heading(doc, "十、实验总结", 1)
    add_text(doc, "本项目最重要的实验结论有三点。第一，结构化数据和确定性规则对 3D 质量门禁非常有效，Rule baseline 在受控测试中达到 95.71% 缺陷 Macro-F1。第二，Qwen2.5-VL 在未经任务适配时对缺陷集合和 Schema 的理解不稳定，但 B4 QLoRA 显著改善了质量判断和结构化缺陷输出。第三，Hybrid 不追求让 VLM 在所有指标上超过规则，而是把模型解释能力和规则安全边界放在同一条审计链中，在冲突时主动复核。", after=8)
    add_text(doc, "因此，这个项目的研究价值不只是一张分数表，而是把数据构造、变量控制、模型适配、评测、错误分析和工程交付串起来，并且明确说明哪些结论可以成立、哪些结论暂时不能成立。", after=8)
    add_heading(doc, "十一、实习总结：我的具体贡献", 1)
    for item in [
        "独立定义了从“3D 资产质检工具”到“面向多图与结构化元数据融合的 VLM 推理研究”的研究问题与边界。",
        "完成 Blender 多视角、UV/法线诊断和六类缺陷注入的数据生成管线，并建立 scene-level split 防止数据泄漏。",
        "完成 Qwen2.5-VL-3B 的推理接口、4-bit QLoRA/SFT 训练、offload 推理和结构化 JSON 评测。",
        "建立 B0-B4 输入消融、三随机种子稳定性、Rule/VLM/Hybrid 同输入对照和错误案例分析。",
        "将研究代码包装为可上传 .blend 的本地 Demo，加入批处理、重试、耗时统计、人工复核路由和 HTML 审计输出。",
        "完成实验报告、论文式中英文草稿、PPT、简历表述和公开 GitHub 仓库，形成可解释、可复现、可展示的项目成果。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "实习收获", "我最大的收获不是“训练了一个 VLM”，而是学会了如何把一个模糊的 AI+3D 想法压缩成可验证的研究问题，再把实验结果转化为有边界、有审计、有工程入口的系统原型。", fill="EEF6FF", label_color=BLUE)
    add_heading(doc, "附录：关键文件与复现入口", 1)
    add_table(doc, ["类别", "文件 / 入口", "作用"], [
        ("研究结果", "reports/phase2_results_blender_v5.md", "600 样本多资产主实验、三种子稳定性与错误分析"),
        ("外部验证", "reports/external_validation_v1.md", "28 个受控 .blend 输入、批处理和三策略对比"),
        ("数据生成", "blender/generate_scene_dataset.py", "Blender 多视角、UV/法线和 metadata 生成"),
        ("批处理", "scripts/run_external_batch.py", "隔离目录、失败重试、日志、P50/P95"),
        ("模式对比", "scripts/compare_modes.py", "Rule-only、VLM-only、Hybrid 同输入评测"),
        ("Demo", "demo/app.py", "上传 .blend、运行审核、导出 HTML 审计"),
        ("公开证据", "reports/demo_reports/", "脱敏的 Hybrid 复核与 clean PASS HTML 报告"),
    ], [1.10, 2.55, 2.85], font_size=8.9)
    add_text(doc, "文档口径说明：主 benchmark 和外部 fixture 都是受控 Blender 数据；项目当前不宣称客户生产数据上的准确率，也不宣称跨医学、自动驾驶等领域泛化。", color=MUTED, size=9.5, italic=True, before=6, after=0)

    doc.core_properties.title = "面向多图与结构化元数据融合的视觉语言模型推理研究 - 实习项目总结报告"
    doc.core_properties.subject = "3D 资产多模态质量检验系统研究与工程实现"
    doc.core_properties.author = ""
    doc.core_properties.comments = ""
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
